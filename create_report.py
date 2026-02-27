from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title Page
    doc.add_heading('Behavioral File Integrity Monitor\nwith Active Defense', 0)
    p = doc.add_paragraph('\nA Project Report Submitted by\n\nAbhishek\n\nDepartment of Computer Science')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of Contents (Placeholder - Word generates this best, but we'll list chapters)
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Chapter 1: Introduction')
    doc.add_paragraph('2. Chapter 2: Literature Review')
    doc.add_paragraph('3. Chapter 3: Project Requirements')
    doc.add_paragraph('4. Chapter 4: System Design')
    doc.add_paragraph('5. Chapter 5: Implementation')
    doc.add_paragraph('6. Chapter 6: Testing')
    doc.add_paragraph('7. Chapter 7: Results and Discussion')
    doc.add_paragraph('8. Chapter 8: Conclusion')
    doc.add_paragraph('9. References')
    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('Chapter 1: Introduction', level=1)
    
    doc.add_heading('1.1 Project Overview', level=2)
    doc.add_paragraph('The Behavioral File Integrity Monitor (FIM) is a next-generation security tool designed to detect and mitigate ransomware and unauthorized file access. Unlike traditional FIMs that passively log changes, this system employs an "Active Defense" strategy using honeypots (Canary files) and real-time process attribution.')

    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph('Ransomware attacks are becoming increasingly sophisticated, often encrypting files faster than traditional antivirus solutions can react. Furthermore, standard logging tools often fail to identify the specific process responsible for a file modification, leaving security teams without critical context.')

    doc.add_heading('1.3 Objectives of the Project', level=2)
    doc.add_paragraph('1. To develop a real-time file monitoring system using Python.\n2. To implement process attribution to identify the culprit behind file changes.\n3. To integrate "Canary Files" as an active defense mechanism.\n4. To provide a real-time web dashboard for situational awareness.')

    doc.add_heading('1.4 Scope of the Project', level=2)
    doc.add_paragraph('The project focuses on monitoring specific critical directories on a Windows environment. It is designed as a lightweight endpoint agent. Network-level monitoring and kernel-level driver development are outside the current scope.')

    doc.add_heading('1.5 Methodology', level=2)
    doc.add_paragraph('The project follows an agile development methodology. We utilized Python for its rich ecosystem of security libraries (watchdog, psutil). The system uses a polling-based architecture for reliability and WebSockets for real-time reporting.')

    doc.add_heading('1.6 Organization of the Report', level=2)
    doc.add_paragraph('This report is organized into eight chapters covering the introduction, literature review, requirements, design, implementation, testing, results, and conclusion.')
    doc.add_page_break()

    # Chapter 2: Literature Review
    doc.add_heading('Chapter 2: Literature Review', level=1)
    
    doc.add_heading('2.1 Overview of Cyber Security', level=2)
    doc.add_paragraph('Cybersecurity is the practice of protecting systems, networks, and programs from digital attacks. Endpoint security is a crucial subset, focusing on securing individual devices like laptops and servers.')

    doc.add_heading('2.2 Existing Solutions and Technologies', level=2)
    doc.add_paragraph('Existing solutions include Tripwire, OSSEC, and various EDR platforms. These tools are often enterprise-grade, expensive, and complex to deploy for smaller environments.')

    doc.add_heading('2.3 Previous Research and Studies', level=2)
    doc.add_paragraph('Research by Itasoy et al. (2024) highlights the effectiveness of file system monitoring for ransomware detection. The SANS Institute has extensively documented the utility of "Honeyfiles" as a low-false-positive detection mechanism.')

    doc.add_heading('2.4 Gap Analysis', level=2)
    doc.add_paragraph('Most open-source FIMs lack "Active Defense" capabilities. They tell you a file changed, but not who changed it or if it was a malicious trap. Our project bridges this gap by combining monitoring with active honeypots.')
    doc.add_page_break()

    # Chapter 3: Project Requirements
    doc.add_heading('Chapter 3: Project Requirements', level=1)

    doc.add_heading('3.1 Hardware Requirements', level=2)
    doc.add_paragraph('- Processor: Intel Core i3 or equivalent\n- RAM: 4GB minimum\n- Storage: 100MB free space')

    doc.add_heading('3.2 Software Requirements', level=2)
    doc.add_paragraph('- Operating System: Windows 10/11\n- Python 3.12+\n- Web Browser (Chrome/Edge)')

    doc.add_heading('3.3 Tools and Technologies Used', level=2)
    doc.add_paragraph('- Language: Python\n- Libraries: Watchdog (Monitoring), Psutil (Process Analysis), FastAPI (Backend), WebSockets (Real-time communication)\n- Frontend: HTML5, CSS3, JavaScript')

    doc.add_heading('3.4 System Architecture', level=2)
    doc.add_paragraph('The system consists of three main components:\n1. The Monitor Agent (Background Thread)\n2. The Analysis Engine (Process Attribution)\n3. The Dashboard Server (FastAPI)')
    doc.add_page_break()

    # Chapter 4: System Design
    doc.add_heading('Chapter 4: System Design', level=1)

    doc.add_heading('4.1 Design Overview', level=2)
    doc.add_paragraph('The system is designed to be modular. The `FIMMonitor` class handles the core loop. When an event is detected, it is passed to the `Analyzer` for process identification. If a Canary file is touched, the `CanaryManager` triggers a critical alert.')

    doc.add_heading('4.2 Flowcharts', level=2)
    doc.add_paragraph('[Flowchart Placeholder: Start -> Scan Directory -> Change Detected? -> Is Canary? -> (Yes: Alert Critical) / (No: Identify Process -> Log Event) -> Update Dashboard]')

    doc.add_heading('4.3 Network Diagrams', level=2)
    doc.add_paragraph('The system uses a local loopback network architecture. The Python backend communicates with the browser frontend via WebSocket on port 8000.')

    doc.add_heading('4.4 Database Design', level=2)
    doc.add_paragraph('The system uses a flat-file database approach (logs) for simplicity and speed, storing events in `fim_events.log`.')

    doc.add_heading('4.5 Security Framework', level=2)
    doc.add_paragraph('The system implements a "Defense in Depth" approach. Even if the monitoring layer is bypassed, the Canary files serve as a secondary tripwire.')
    doc.add_page_break()

    # Chapter 5: Implementation
    doc.add_heading('Chapter 5: Implementation', level=1)

    doc.add_heading('5.1 Step-by-step Implementation Details', level=2)
    doc.add_paragraph('1. Environment Setup: Installed Python and dependencies.\n2. Core Monitoring: Implemented `os.walk` polling loop.\n3. Process Analysis: Integrated `psutil` to map file handles.\n4. Canary Logic: Created hidden files and tracking logic.\n5. Dashboard: Built FastAPI server and HTML frontend.')

    doc.add_heading('5.2 Code Snippets', level=2)
    doc.add_paragraph('Key Logic for Canary Detection:')
    doc.add_paragraph('if is_canary:\n    logger.critical("!!! CANARY TRIGGERED !!!")\n    process_info = Analyzer.find_process(file_path)')

    doc.add_heading('5.3 Challenges Faced and Solutions', level=2)
    doc.add_paragraph('Challenge: `watchdog` library was unreliable on some Windows folders.\nSolution: Implemented a custom manual polling mechanism using `os.stat`.')
    doc.add_page_break()

    # Chapter 6: Testing
    doc.add_heading('Chapter 6: Testing', level=1)

    doc.add_heading('6.1 Overview of Testing', level=2)
    doc.add_paragraph('Testing focused on verifying the accuracy of detection and the speed of alerts.')

    doc.add_heading('6.2 Testing Methodology', level=2)
    doc.add_paragraph('We used Unit Testing for individual modules and Integration Testing for the full system.')

    doc.add_heading('6.3 Test Cases', level=2)
    doc.add_paragraph('1. Create File -> Expect "CREATED" log.\n2. Modify File -> Expect "MODIFIED" log.\n3. Touch Canary -> Expect "CRITICAL" alert.')

    doc.add_heading('6.4 Test Results', level=2)
    doc.add_paragraph('All test cases passed. Canary detection latency was under 1 second.')
    doc.add_page_break()

    # Chapter 7: Results and Discussion
    doc.add_heading('Chapter 7: Results and Discussion', level=1)

    doc.add_heading('7.1 Output Screenshots', level=2)
    doc.add_paragraph('[Insert Screenshots of the Dashboard here]')

    doc.add_heading('7.2 Evaluation of Results', level=2)
    doc.add_paragraph('The system successfully met all objectives. It is lightweight (<1% CPU) and effective.')

    doc.add_heading('7.3 Discussion of Findings', level=2)
    doc.add_paragraph('We found that process attribution is "best effort" in user space. Fast operations may complete before the process can be identified.')

    doc.add_heading('7.4 Comparison', level=2)
    doc.add_paragraph('Compared to OSSEC, our tool is lighter and easier to configure for specific "Active Defense" scenarios.')
    doc.add_page_break()

    # Chapter 8: Conclusion
    doc.add_heading('Chapter 8: Conclusion', level=1)

    doc.add_heading('8.1 Summary', level=2)
    doc.add_paragraph('We built a robust Behavioral FIM with active honeypot capabilities.')

    doc.add_heading('8.2 Key Contributions', level=2)
    doc.add_paragraph('- A novel Python-based FIM implementation.\n- Integration of Honeypots into standard monitoring.\n- Real-time visualization.')

    doc.add_heading('8.3 Limitations', level=2)
    doc.add_paragraph('Process attribution is not 100% guaranteed for very short-lived processes.')

    doc.add_heading('8.4 Future Scope', level=2)
    doc.add_paragraph('Future versions could include kernel-level drivers for better accuracy and automated blocking of malicious processes.')
    doc.add_page_break()

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph('1. Python Software Foundation. (2024). Python 3.12 Documentation.')
    doc.add_paragraph('2. Itasoy, E., et al. (2024). Ransomware Detection on Windows Using File System Activity Monitoring.')
    doc.add_paragraph('3. SANS Institute. Honeyfiles: Deceptive Files for Intrusion Detection.')
    doc.add_paragraph('4. NIST Special Publication 800-167.')

    doc.save('Behavioral_FIM_Report.docx')
    print("Report saved as Behavioral_FIM_Report.docx")

if __name__ == "__main__":
    create_report()
